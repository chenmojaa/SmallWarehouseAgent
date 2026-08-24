"""Parse .docx -> Markdown text via python-docx, preserving merged cells.

Walks the underlying <w:tbl> / <w:tr> / <w:tc> XML directly so we can read
gridSpan / vMerge attributes reliably (python-docx's _Cell pooling makes
id()-based dedup unreliable).
"""
from __future__ import annotations

from docx import Document
from docx.oxml.ns import qn


def _md_escape(s: str) -> str:
    return s.replace("|", "\\|").replace("\n", " ").replace("\r", " ")


def _tc_grid_span(tc) -> int:
    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is None:
        return 1
    gs = tcPr.find(qn("w:gridSpan"))
    if gs is None:
        return 1
    try:
        return max(1, int(gs.get(qn("w:val"))))
    except (TypeError, ValueError):
        return 1


def _tc_vmerge_state(tc):
    """Return 'restart', 'continue', or None based on w:vMerge."""
    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is None:
        return None
    vm = tcPr.find(qn("w:vMerge"))
    if vm is None:
        return None
    val = vm.get(qn("w:val"))
    return "restart" if val == "restart" else "continue"


def _tc_text(tc) -> str:
    return "".join(t.text or "" for t in tc.iter(qn("w:t"))).strip()


def _grid_to_markdown(rows: list[list[str]]) -> str:
    if not rows:
        return ""
    n_cols = max((len(r) for r in rows), default=0)
    rows = [r + [""] * (n_cols - len(r)) for r in rows]
    while rows and not any(c.strip() for c in rows[-1]):
        rows.pop()
    if not rows:
        return ""
    widths = [0] * n_cols
    for r in rows:
        for i, c in enumerate(r):
            widths[i] = max(widths[i], len(_md_escape(c)))

    def fmt_row(r):
        return "| " + " | ".join(_md_escape(c).ljust(widths[i]) for i, c in enumerate(r)) + " |"

    header = rows[0]
    sep = "| " + " | ".join("-" * widths[i] for i in range(n_cols)) + " |"
    body = "\n".join(fmt_row(r) for r in rows[1:])
    out = fmt_row(header) + "\n" + sep
    if body:
        out += "\n" + body
    return out


def _table_to_markdown(table) -> str:
    """Convert a docx table to Markdown by walking raw <w:tc> elements."""
    tbl = table._tbl
    rows_out: list[list[str]] = []
    # Track vMerge column -> restart text so continuation cells can be blanked.
    vmerge_anchor_text: dict[int, str] = {}

    for tr in tbl.iter(qn("w:tr")):
        cells: list[str] = []
        col_idx = 0
        for tc in tr.findall(qn("w:tc")):
            grid_span = _tc_grid_span(tc)
            vmerge = _tc_vmerge_state(tc)
            if vmerge == "continue":
                # Continuation cell: emit the anchor's text (visual = single merged value).
                anchor_text = vmerge_anchor_text.get(col_idx, "")
                cells.append(anchor_text)
                for _ in range(grid_span - 1):
                    cells.append("")
                col_idx += grid_span
                continue
            text = _tc_text(tc)
            cells.append(text)
            for _ in range(grid_span - 1):
                cells.append("")
            if vmerge == "restart":
                vmerge_anchor_text[col_idx] = text
            col_idx += grid_span
        rows_out.append(cells)

    return _grid_to_markdown(rows_out)


def parse_docx(path: str) -> dict:
    doc = Document(path)
    parts = []
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text)
    for table in doc.tables:
        md = _table_to_markdown(table)
        if md:
            parts.append(md)
    content = "\n\n".join(parts).strip()
    title = ""
    if doc.core_properties and doc.core_properties.title:
        title = str(doc.core_properties.title)[:200]
    if not title:
        for p in doc.paragraphs:
            if p.text.strip():
                title = p.text.strip()[:80]
                break
    if not title:
        title = "Untitled Document"
    return {"title": title, "content": content or "(empty document)"}
